import tempfile
import os
from markdown2 import markdown
import pdfkit  
import pandas as pd
from io import StringIO
import json
from docx import Document

def convert_markdown_to_format(md_content: str, fmt: str) -> str:
    tmp_dir = tempfile.gettempdir()

    if fmt == "html":
        html_content = markdown(md_content, extras=["tables"])
        styled_html = f"""
        <html>
        <head>
        <meta charset="utf-8">
        <style>
            body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            padding: 20px;
            }}
            table {{
            border-collapse: collapse;
            width: 100%;
            margin-top: 20px;
            }}
            th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: center;
            }}
            th {{
            background-color: #f4f4f4;
            font-weight: bold;
            }}
            tr:nth-child(even) {{
            background-color: #fafafa;
            }}
        </style>
        </head>
        <body>
        {html_content}
        </body>
        </html>
        """
        output_path = os.path.join(tmp_dir, "output.html")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(styled_html)

    elif fmt == "pdf":
        html_content = markdown(md_content, extras=["tables"])
        styled_html = f"""
        <html>
        <head>
        <meta charset="utf-8">
        <style>
            body {{
            font-family: "Apple SD Gothic Neo", "Nanum Gothic", sans-serif;
            line-height: 1.6;
            padding: 20px;
            font-size: 12pt;
            }}
            table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
            font-size: 10pt;
            }}
            th, td {{
            border: 1px solid #444;
            padding: 8px 12px;
            text-align: center;
            }}
            th {{
            background-color: #f2f2f2;
            font-weight: bold;
            }}
            tr:nth-child(even) {{
            background-color: #fafafa;
            }}
        </style>
        </head>
        <body>
        {html_content}
        </body>
        </html>
        """

        html_path = os.path.join(tmp_dir, "temp.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(styled_html)

        output_path = os.path.join(tmp_dir, "output.pdf")
        pdfkit.from_file(html_path, output_path)

    elif fmt == "csv":
        lines = [line for line in md_content.splitlines() if "|" in line and "---" not in line]
        if len(lines) < 2:
            raise ValueError("유효한 마크다운 테이블이 아닙니다.")
        table = "\n".join(lines)
        df = pd.read_csv(StringIO(table), sep="|", engine="python", skipinitialspace=True)
        df = df.dropna(axis=1, how="all")  # 빈 열 제거
        output_path = os.path.join(tmp_dir, "output.csv")
        df.to_csv(output_path, index=False)

    elif fmt == "xlsx":
        lines = [line for line in md_content.splitlines() if "|" in line and "---" not in line]
        if len(lines) < 2:
            raise ValueError("유효한 마크다운 테이블이 아닙니다.")
        table = "\n".join(lines)
        df = pd.read_csv(StringIO(table), sep="|", engine="python", skipinitialspace=True)
        df = df.dropna(axis=1, how="all")
        output_path = os.path.join(tmp_dir, "output.xlsx")
        df.to_excel(output_path, index=False)

    elif fmt == "json":
        lines = [line for line in md_content.splitlines() if "|" in line and "---" not in line]
        if len(lines) < 2:
            raise ValueError("유효한 마크다운 테이블이 아닙니다.")
        table = "\n".join(lines)
        df = pd.read_csv(StringIO(table), sep="|", engine="python", skipinitialspace=True)
        df = df.dropna(axis=1, how="all")
        output_path = os.path.join(tmp_dir, "output.json")
        df.to_json(output_path, orient="records", force_ascii=False, indent=2)

    elif fmt == "docx":
        document = Document()
        lines = [line.strip() for line in md_content.splitlines() if "|" in line and "---" not in line]

        if lines:
            # 첫 줄은 헤더
            header = [h.strip() for h in lines[0].split("|") if h.strip()]
            rows = [[c.strip() for c in row.split("|") if c.strip()] for row in lines[1:]]

            # 표 생성
            table = document.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"

            # 헤더 채우기
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(header):
                hdr_cells[i].text = h

            # 데이터 채우기
            for row in rows:
                row_cells = table.add_row().cells
                for i, c in enumerate(row):
                    row_cells[i].text = c

            document.add_paragraph("\n")  # 표 이후 여백

        # 표 외 텍스트도 출력
        for line in md_content.splitlines():
            if "|" not in line:
                document.add_paragraph(line)

        output_path = os.path.join(tmp_dir, "output.docx")
        document.save(output_path)

    else:
        raise ValueError("지원하지 않는 포맷")

    return output_path